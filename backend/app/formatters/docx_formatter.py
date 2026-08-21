import re
import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, List


class DocxFormatter:
    """
    Fills DOCX templates:
      1. fill_placeholders: for templates with {{placeholder}} syntax
      2. freeform_fill: for templates with existing styles to mirror
      3. create_structured_cv: generates a professional, clean DOCX from extracted CV JSON
    """

    def fill_placeholders(self, template_path: str, mapping: Dict[str, str], output_path: str) -> str:
        """
        Replace all {{placeholder}} occurrences in template with mapped values.
        Works on paragraphs, tables, headers, and footers.
        """
        doc = docx.Document(template_path)

        # Process main body paragraphs
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, mapping)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, mapping)

        # Process headers and footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                self._replace_in_paragraph(para, mapping)
            for para in section.footer.paragraphs:
                self._replace_in_paragraph(para, mapping)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self._replace_in_paragraph(para, mapping)

        doc.save(output_path)
        return output_path

    def _replace_in_paragraph(self, paragraph, mapping: Dict[str, str]):
        if not paragraph.text:
            return

        full_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
        new_text = full_text

        for placeholder, value in mapping.items():
            if placeholder in new_text:
                replacement = str(value) if value is not None else ""
                new_text = new_text.replace(placeholder, replacement)

        if new_text == full_text:
            return

        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(new_text)

    def create_structured_cv(self, cv_data: Dict[str, Any], output_path: str) -> str:
        """Create a professional, beautifully styled CV document from structured CV data."""
        doc = docx.Document()

        # Set page margins
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        # Header - Candidate Name & Title
        name = cv_data.get("full_name") or f"{cv_data.get('first_name', '')} {cv_data.get('last_name', '')}".strip() or "Candidate"
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(name)
        name_run.font.size = Pt(20)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(30, 41, 59) # Slate 800

        job_title = cv_data.get("job_title")
        if job_title:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(job_title)
            title_run.font.size = Pt(12)
            title_run.font.italic = True
            title_run.font.color.rgb = RGBColor(79, 70, 229) # Indigo 600

        # Contact Info Line
        contact_items = []
        if cv_data.get("location"): contact_items.append(cv_data["location"])
        elif cv_data.get("city") or cv_data.get("country"):
            loc = ", ".join(filter(None, [cv_data.get("city"), cv_data.get("country")]))
            if loc: contact_items.append(loc)
        if cv_data.get("email"): contact_items.append(cv_data["email"])
        if cv_data.get("phone"): contact_items.append(cv_data["phone"])
        if cv_data.get("linkedin"): contact_items.append(cv_data["linkedin"])
        if cv_data.get("github"): contact_items.append(cv_data["github"])

        if contact_items:
            contact_para = doc.add_paragraph()
            contact_run = contact_para.add_run(" • ".join(contact_items))
            contact_run.font.size = Pt(9.5)
            contact_run.font.color.rgb = RGBColor(100, 116, 139)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # Helper to add section headers
        def add_section_header(title: str):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(title.upper())
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = RGBColor(79, 70, 229)

        # Professional Summary
        summary = cv_data.get("professional_summary")
        if summary:
            add_section_header("Professional Summary")
            p = doc.add_paragraph()
            r = p.add_run(summary)
            r.font.size = Pt(10)
            p.paragraph_format.line_spacing = 1.15

        # Technical Skills
        skills = cv_data.get("technical_skills") or cv_data.get("skills")
        if skills:
            add_section_header("Technical Skills")
            if isinstance(skills, dict):
                for category, skill_list in skills.items():
                    if skill_list:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(2)
                        cat_run = p.add_run(f"{category.replace('_', ' ').title()}: ")
                        cat_run.font.bold = True
                        cat_run.font.size = Pt(9.5)
                        val_run = p.add_run(", ".join(skill_list) if isinstance(skill_list, list) else str(skill_list))
                        val_run.font.size = Pt(9.5)
            elif isinstance(skills, list):
                p = doc.add_paragraph()
                r = p.add_run(", ".join(skills))
                r.font.size = Pt(9.5)

        # Professional Experience
        experiences = cv_data.get("experience") or []
        if experiences:
            add_section_header("Professional Experience")
            for exp in experiences:
                company = exp.get("company", "")
                title = exp.get("title", "")
                dates = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')}".strip(" –")
                loc = exp.get("location", "")

                role_para = doc.add_paragraph()
                role_para.paragraph_format.space_before = Pt(6)
                role_para.paragraph_format.space_after = Pt(2)

                t_run = role_para.add_run(f"{title} | {company}")
                t_run.font.bold = True
                t_run.font.size = Pt(10.5)

                if dates or loc:
                    meta = " | ".join(filter(None, [dates, loc]))
                    d_run = role_para.add_run(f"  ({meta})")
                    d_run.font.italic = True
                    d_run.font.size = Pt(9.5)
                    d_run.font.color.rgb = RGBColor(100, 116, 139)

                bullets = exp.get("bullets") or exp.get("responsibilities") or []
                for b in bullets:
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_after = Pt(2)
                    br = bp.add_run(str(b).lstrip("•-* "))
                    br.font.size = Pt(9.5)
                    bp.paragraph_format.line_spacing = 1.15

        # Key Projects
        projects = cv_data.get("projects") or []
        if projects:
            add_section_header("Key Projects & Deliverables")
            for proj in projects:
                p_title = proj.get("title", "")
                tech = proj.get("tech_stack", "")
                p_para = doc.add_paragraph()
                p_para.paragraph_format.space_before = Pt(4)
                p_para.paragraph_format.space_after = Pt(1)
                tr = p_para.add_run(p_title)
                tr.font.bold = True
                tr.font.size = Pt(10)
                if tech:
                    tech_run = p_para.add_run(f" — Tech: {tech}")
                    tech_run.font.italic = True
                    tech_run.font.size = Pt(9)
                    tech_run.font.color.rgb = RGBColor(100, 116, 139)

                p_bullets = proj.get("bullets") or []
                for b in p_bullets:
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_after = Pt(2)
                    br = bp.add_run(str(b).lstrip("•-* "))
                    br.font.size = Pt(9.5)

        # Education
        education = cv_data.get("education") or []
        if education:
            add_section_header("Education")
            for edu in education:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                degree = edu.get("degree", "")
                inst = edu.get("institution", "")
                yr = edu.get("graduation_year", "")
                r1 = p.add_run(f"{degree} — {inst}")
                r1.font.bold = True
                r1.font.size = Pt(9.5)
                if yr:
                    r2 = p.add_run(f" ({yr})")
                    r2.font.italic = True
                    r2.font.size = Pt(9)

        # Certifications
        certs = cv_data.get("certifications") or []
        if certs:
            add_section_header("Certifications")
            for c in certs:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                c_title = c.get("title") or c.get("name", "")
                c_issuer = c.get("issuer", "")
                c_yr = c.get("year") or c.get("date", "")
                desc = " • ".join(filter(None, [c_title, c_issuer, c_yr]))
                r = p.add_run(desc)
                r.font.size = Pt(9.5)

        doc.save(output_path)
        return output_path
