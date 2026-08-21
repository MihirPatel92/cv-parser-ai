import docx
import re
from typing import List


def extract_text_from_docx(file_path: str) -> str:
    """Extract all text from paragraphs, tables, headers, and footers of a DOCX file."""
    try:
        doc = docx.Document(file_path)
        full_text = []

        # Body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            row_texts.append(para.text.strip())
                if row_texts:
                    full_text.append(" | ".join(row_texts))

        # Headers & Footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            for para in section.footer.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

        return "\n".join(full_text)
    except Exception as e:
        raise Exception(f"Failed to parse DOCX: {e}")


def extract_placeholders_from_docx(file_path: str) -> List[str]:
    """Extract all unique {{placeholder}} occurrences from doc paragraphs, tables, headers, footers."""
    try:
        doc = docx.Document(file_path)
        raw_text_chunks = []

        # Paragraphs
        for para in doc.paragraphs:
            raw_text_chunks.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        raw_text_chunks.append(para.text)

        # Headers & Footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                raw_text_chunks.append(para.text)
            for para in section.footer.paragraphs:
                raw_text_chunks.append(para.text)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            raw_text_chunks.append(para.text)

        combined = " ".join(raw_text_chunks)
        # Match patterns like {{First_Name}}, {{ JOB_TITLE }}, {{Role_1_Company_Name}}, etc.
        placeholders = re.findall(r"\{\{[^}]+\}\}", combined)
        # Clean and deduplicate while preserving stable order
        seen = set()
        unique_placeholders = []
        for p in placeholders:
            cleaned = p.strip()
            if cleaned not in seen:
                seen.add(cleaned)
                unique_placeholders.append(cleaned)

        return unique_placeholders
    except Exception as e:
        print(f"Error extracting placeholders from DOCX: {e}")
        return []
